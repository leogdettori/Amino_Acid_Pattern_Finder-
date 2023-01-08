## ====================================================================================================================================== ## 
##                                                                                                                                        ##
##                            First, loading some Functions required for Amino_Acid_Pattern_Finder_Extended                               ##
##                               Find Amino_Acid_Pattern_Finder_Extended at the end of this sections!!!                                   ##
##                                                                                                                                        ##
## ====================================================================================================================================== ## 

"""
Created on Fri Apr 19 17:18:06 2019
Modified on Fri Jan 08 2021 by LD

@author: pappc, LD
"""

import numpy as np
import pandas as pd
import plotly.graph_objs as go
import plotly.figure_factory as ff
import matplotlib.pyplot as plt

## ======================================================================== ###

#Find the non-common values in two lists
def Diff(list1, list2):
    return (list(list(set(list1)-set(list2)) + list(set(list2)-set(list1))))

## ======================================================================== ###

def occurrences(string, sub):
    count = start = 0
    while True:
        start = string.find(sub, start) + 1
        if start > 0:
            count+=1
        else:
            return count
           
## ======================================================================== ###


def read_fasta (in_fasta):
    with open(in_fasta) as in_data_file:
        lines = in_data_file.readlines()
    seq = ''
    for line in lines:
        if line.startswith ('>') == False:
            seq = r''.join((seq, line.strip()))
    return seq

## ======================================================================== ###


def show_values_on_bars(axs):
    def _show_on_single_plot(ax):        
        for p in ax.patches:
            _x = p.get_x() + p.get_width() / 2
            _y = p.get_y() + p.get_height() + 0.1
            value = '{:.1f}'.format(p.get_height())
            ax.text(_x, _y, value, ha="center", size=9) 

    if isinstance(axs, np.ndarray):
        for idx, ax in np.ndenumerate(axs):
            _show_on_single_plot(ax)
    else:
        _show_on_single_plot(axs)
## ======================================================================== ###
#transform the sequence using a reduced amino acid alphabet
def transform_sequence(seq, groups, group_names):
    group_list = []
    seq_new = seq
    for i in range(len(groups)):
        group_list.append(groups[i])
        for a in group_list[i]:
            seq_new = seq_new.replace(a, group_names[i])
    return seq_new
    
## ======================================================================== ###
def create_aa_df(seq, aa_alphabet):
    #dataframe creation
    aa_counts = []
    for i in aa_alphabet:
        count = occurrences(seq, i)
        aa_counts.append(count)
    df = pd.DataFrame(aa_alphabet, columns = ['aa'])
    df['count'] = aa_counts
    aa_freq = []
    for i in aa_counts:
        value = (i/len(seq))*100
        #value = round(value, 2)    
        aa_freq.append(value)
    df['freq'] = aa_freq
    return df

## ======================================================================== ###
#first plot
def freq_plot(dataframe, x, y, Name = None, sort = False):
    if sort:
        dataframe = dataframe.sort_values(y, ascending = False)
    data = [go.Bar(
            x=dataframe[x],
            y=dataframe[y],
            text=dataframe[y],
            textposition = 'outside',
            textfont=dict(
                    size=18),
            marker=dict(
                color='rgb(158,202,250)',
                line=dict(
                    color='rgb(8,48,107)',
                    width=1.5)            
            )
        )]             
    layout = go.Layout(
            title='Amino acid frequencies',
            xaxis={'title':'Amino acids'},
            yaxis=dict(
                    title = 'Frequency'
                    )
            )
    if Name:
        layout.update(dict(title = 'Amino acid frequencies in {}'.format(Name)))
    figure = go.Figure(data=data, layout=layout)
    return figure
## ======================================================================== ###    
#dipeptide plot
#takes 3 argument: 1.)the sequence which MUST be transformed before using a reduced amino acid alphabet
# 2.) the actual alphabet, which should be a list of letters/numbers
# 3.) the 'groups', which should be a list of lists, containing the grouped amino acids

def dipep_plot(seq, alphabet, groups = None):
    dipeptides = [a + b for a in alphabet for b in alphabet]
    data_di = []
    for k in dipeptides:
        motif_count = occurrences(seq, k)
        data_di.append(motif_count)
    di_df = pd.DataFrame(dipeptides, columns = ['dipeptide'])
    dipep_freq = []
    for i in data_di:
        if i != 0:
            frequency = (i/len(seq))*100
            frequency = round(frequency)    
            dipep_freq.append(frequency)
        else:
            dipep_freq.append(i)
    di_df['freq'] = dipep_freq
    #create array for heatmap
    num_aa = len(alphabet)
    array = np.array(di_df['freq']).reshape(num_aa, num_aa)
    #use plotly to create heatmap
    fig = ff.create_annotated_heatmap(array,
                                     x=alphabet,
                                     y=alphabet,
                                     colorscale = 'Viridis')
    fig.layout.update(
            go.Layout(
                    autosize=False,
                    width=500,
                    height=500,
                    xaxis = dict(
                            side = 'bottom',
                            type= 'category'
                            ),
                    yaxis = dict(
                            type='category'),
                    title = 'Dipeptide frequencies (%) using {} letters'.format(num_aa)
                    )
            )
    return fig

## ======================================================================== TOP - LD  1/28/2021    
#tripeptide plot values for subsequent creation of txt file table
#takes 3 argument: 1.)the sequence which MUST be transformed before using a reduced amino acid alphabet
# 2.) the actual alphabet, which should be a list of letters/numbers
# 3.) the 'groups', which should be a list of lists, containing the grouped amino acids

def tripep_plot_values(seq, alphabet, groups = None):
    tripeptides = [a + b + c for a in alphabet for b in alphabet for c in alphabet]
    data_tri = []
    for k in tripeptides:
        motif_count = occurrences(seq, k)
        data_tri.append(motif_count)
    tri_df = pd.DataFrame(tripeptides, columns = ['tripeptide'])
    tripep_freq = []
    for i in data_tri:
        if i != 0:
            frequency = (i/(len(seq)-2))*100    #changed on 08/09/2021 by LD, DT
            #frequency = (i/len(seq))*100
            #frequency = round(frequency)    
            tripep_freq.append(frequency)
        else:
            tripep_freq.append(i)
    tri_df['freq'] = tripep_freq

    return tri_df


## ======================================================================== TOP - LD  1/8/2021
#dipeptide plot values for subsequent creation of txt file table
#takes 3 argument: 1.)the sequence which MUST be transformed before using a reduced amino acid alphabet
# 2.) the actual alphabet, which should be a list of letters/numbers
# 3.) the 'groups', which should be a list of lists, containing the grouped amino acids

def dipep_plot_values(seq, alphabet, groups = None):
    dipeptides = [a + b for a in alphabet for b in alphabet]
    data_di = []
    for k in dipeptides:
        motif_count = occurrences(seq, k)
        data_di.append(motif_count)
    di_df = pd.DataFrame(dipeptides, columns = ['dipeptide'])
    dipep_freq = []
    for i in data_di:
        if i != 0:
            frequency = (i/(len(seq)-1))*100    #changed on 08/09/2021 by LD, DT
            #frequency = (i/len(seq))*100
            #frequency = round(frequency)    
            dipep_freq.append(frequency)
        else:
            dipep_freq.append(i)
    di_df['freq'] = dipep_freq

    return di_df


## ======================================================================== ###  
def pie_chart(OUTPUT, dataframe, x, y, plots = "on", Name = None):
    #Sorting data
    dataframe = dataframe.sort_values(y, ascending = False)
    #print(dataframe)
    
    #Rearranging data to plot
    i1 = 0
    i2 = 0
    data = {'labels':[], 'freq':[]}
    others_freq= 100
    for element1 in dataframe[x]:              
        if i1 < 3:
            data['labels'].append(element1)
            i1 = i1 + 1
            
    for element2 in dataframe[y]:              
        if i2 < 3:
            data['freq'].append(element2)
            i2 = i2 + 1
            others_freq = others_freq - element2        
    
    #Adding the "Others" for all the remaining amino acids
    data['labels'].append('Others')
    data['freq'].append(others_freq)
    #print(data['labels'])
    #print(data['freq'])
    
    # Plotting
    #Only creates plots if the user chose to:
    if plots != "off":    
        plt.figure(figsize=(8,5))
        colors = ['gold', 'yellowgreen', 'lightcoral', 'lightskyblue']
        plt.title('Amino acid frequencies in {}'.format(Name),fontsize='x-large')
        fig1 = plt.pie(data['freq'], labels=data['labels'], autopct='%1.1f%%', colors=colors,  textprops={'fontsize':14}, shadow=True, startangle=140)
        plt.legend(loc="upper right",fontsize='medium',bbox_to_anchor=(1.4,1))

        plt.savefig(OUTPUT+'_pie_chart.png')
        plt.close()
    
        #plt.pie(sizes, explode=explode, labels=labels, colors=colors,autopct='%1.1f%%', shadow=True, startangle=140)
        #colors = ['gold', 'yellowgreen', 'lightcoral', 'lightskyblue']
        #explode = (0.1, 0, 0, 0)  # explode 1st slice
        #plt.legend()
        #plt.axis('equal')
        #plt.show()
    
    return data

## ======================================================================== Bottom - LD  1/8/2021



## ====================================================================================================================================== ## 
##                                                                                                                                        ##
##                                          Amino_Acid_Pattern_Finder_Extended starts here!!!                                             ##
##                                                                                                                                        ##
## ====================================================================================================================================== ## 

"""
Created on Mon May 13 13:06:08 2019
Modified on Wed Feb 24 2021 by LD

Script for finding and plotting the frequency of different dipeptide combinations
in selected protein sequences

@authors: Csaba Papp, Leo Dettori

"""


def Amino_Acid_Pattern_Finder_Extended(list_file_location, plots = "on", output_location = "input"):
    import sys, os
    import matplotlib.pyplot as plt
    #from functions_extended_LD_01282021 import * #from previous version
    import plotly.io as pio
    import docx
    from docx.enum.text import WD_COLOR_INDEX

    #Starting message:
    print("Amino_Acid_Pattern_Finder_Extended has started:\n\n")
    
    #Plots message:
    if plots == "off":
        print("Plots are deactivated!")
    
    
    #Switches directory to the input location, which should contain the coded_lists from Coded_List_Generator    
    os.chdir(list_file_location)

    #Formats input location if need be
    if list_file_location.endswith("/"):
        pass
    else:
        list_file_location = list_file_location +"/"

    #Compiles list of files within the input directory
    files = next(os.walk(list_file_location))[2]
    #print(files)

    #Creates variable to keep track of whether we come directly from Coded_List_Generator
    independently = 0
    #Starts very handy counter
    n1 = 0  #Helps to define the amount of groups we will be working with
    n2 = 0  #Keeps track of the group that we are currently working on

    #In case output was not defined
    if output_location == "input":
        output_location = list_file_location
        print("No output defined! Output was set to same destination as input!" + "\n")

        #Checks if we're coming directly from the Coded_List_Generator
        check_location = output_location.split("Coded_Lists")
        if len(check_location) > 1:
            #print("yes")
            #Switches directory to output location
            os.chdir(output_location)        
            #Switches directory to folder above
            path_parent = os.path.dirname(os.getcwd())
            os.chdir(path_parent)
            #Updates output path information
            output_location = str(path_parent).replace('\\', "/")+"/"   
        #In case we didn't come directly from Coded_List_Generator    
        else:
            independently = 1

    #Neat Message
    print("# --------------------------------------------- #")
    
    
    #Prepares to create a master folder to store the database dett files
    if output_location.endswith("/"):
        pass
    else:
        output_location = output_location + "/"
    #Creates the master folder to store database dett files in an orderly manner                 
    os.mkdir(output_location+"Database")
    #Updates output location
    output_database = output_location+"Database/"    
    
    
    #Prepares to create the master folder to store results in an orderly manner
    if output_location.endswith("/"):
        pass
    else:
        output_location = output_location + "/"
    #Creates the master folder to store results in an orderly manner                 
    os.mkdir(output_location+"Amino_Acid_Pattern_Finder_Results")
    #Updates output location
    output_location = output_location+"Amino_Acid_Pattern_Finder_Results/"


    #Starts to loop through the actual Amino_Acid_Pattern_Finder Program    
    for f in files:

    ################################################################################################################################
    ################################################################################################################################
    ################################################################################################################################
    #                                  Start of Original Amino_Acid_Pattern_Finder_Extended Program                                #
    ################################################################################################################################
    ################################################################################################################################
    ################################################################################################################################

        #Here we loop through each protein sequence within a group
        
        #Updates run name for current run
        run_name = str(f).strip(".txt")

        ## ======================================================================== TOP - LD  2/24/2021

        #Neat message        
        print("\n" + 'Running Amino_Acid_Pattern_Finder_Extended for:' +"\n"+ str(f).strip(".txt"))
        #print("\n" + 'Currently processing:')

        #Creates subfolder to store current run
        os.mkdir(output_location+str(f).strip(".txt")+"/")

        #Updates output location for current run
        output_dir = output_location+str(f).strip(".txt")+"/"
        #print(output_dir)

        #Updates subject of current run
        list_file = list_file_location+str(f)
        #print(list_file) 

        ## ======================================================================== Bottom - LD  2/24/2021

        #load original 20 amino acids for the frequency plots
        all_amino_acids = ['I','V','L','F','C','M','A','G','T','W','S','Y','P','H','N','Q', 'E','D','K','R']
        #print(all_amino_acids)

        #load the colors used for the colored text later in the docx later
        values = ['RED', 'GRAY_50', 'GREEN', 'YELLOW', 'PINK', 'BLUE', 'TEAL', 'TURQUOISE', 'VIOLET']
        color_dict = dict(list(enumerate(values)))
        if output_dir[-1] != '/':
            output_dir += '/'
        #open the docx, iterate through all lines in the list_file and execute the functions
        document = docx.Document()

        ## ======================================================================== TOP - LD  2/25/2021

        #Creates the text files to save the results in table form  --- Modified on 1/28/2021
        f2 = open(str(output_dir)+"Table.txt", "a")
        f2.close()
        f3 = open(str(output_dir)+"Extended_Table.txt", "a")
        f3.close()
        #Establishes table_counter to later create the header of the tables
        table_counter = 0
        table_counter_extended = 0
        
        
        #Creates the text files to save the results in averaged top frequencies format
        #For Dipeptide amino acid
        f4 = open(str(output_dir)+"Top_dipeptide_aa_average.txt", "a")
        f4.write("Most Abundant Dipeptide as Amino Acid in Average (%):\n")
        f4.close()
        #For Dipeptide clustered
        f5 = open(str(output_dir)+"Top_dipeptide_clustered_average.txt", "a")
        f5.write("Most Abundant Dipeptide Clustered in Average (%):\n")
        f5.close()
        #For Tripeptide amino acid
        f7 = open(str(output_dir)+"Top_tripeptide_aa_average.txt", "a")
        f7.write("Most Abundant Tripeptide as Amino Acid in Average (%):\n")
        f7.close()
        #For Tripeptide clustered
        f8 = open(str(output_dir)+"Top_tripeptide_clustered_average.txt", "a")
        f8.write("Most Abundant Tripeptide Clustered in Average (%):\n")
        f8.close()
        #For single amino acid
        f9 = open(str(output_dir)+"Top_peptide_aa_average.txt", "a")
        f9.write("Most Abundant Amino Acids in Average (%):\n")
        f9.close()
        #For single peptide clustered
        f10 = open(str(output_dir)+"Top_peptide_clustered_average.txt", "a")
        f10.write("Most Abundant Peptide Clustered in Average (%):\n")
        f10.close()
        
       

        #Creates a neat counter to keep track of how many protein sequences exist within this group
        counter = 0  #this counter also helps with the averaged top frequencies calculation

        ## ======================================================================== Bottom - LD  2/25/2021

        #open file, read in and do basic formatting on data
        with open(list_file, "r") as f:
            data = f.readlines()
            lines = []
            for a in data:
                lines.append(a.strip())
            data = list(filter(None, lines))
        #assign the data in the different columns to the respective variables
        for count, line in enumerate(data):
            #split the contents of the line into separate variables
            name, file, group_codes, groups = line.strip().split('\t')
            #evaluation of certain variables is necessary, so they are stored as LISTS instead of strings
            groups = eval(groups)
            #in this step we will identify the amino acids that WILL REMAIN UNCHANGED
            aa_list = []
            group_codes = eval(group_codes)
            #create a single list from all of our groups for easier comparison
            for group in groups:
                for aa in group:
                    aa_list.append(aa.upper())
            #the Diff function will identify the elements that are not present in both lists
            #i.e. the amino acids we will not be changing
            diff = Diff(aa_list, all_amino_acids)
            #by combining the unchanged amino acids with our coded groups, we have the complete list of letter we will be using
            amino_acids = group_codes + diff
            #read in fasta file, lines starting with ">" will be ignored
            seq = read_fasta(file)
            #set output directory
            output = os.path.join(output_dir, str(count) + '_' + name)
            #Create amino acid frequency table and plots
            df = create_aa_df(seq, all_amino_acids)
            #Only creates plots if the user chose to:
            if plots != "off":
                freq_plot1 = freq_plot(df, 'aa', 'freq', name)
                pio.write_image(freq_plot1, output + '_freq_plot.png')
                freq_plot2 = freq_plot(df, 'aa', 'freq', name, True)
                pio.write_image(freq_plot2, output + '_sorted_freq_plot.png')


            #Transform the sequence according to the provided groups and create plots
            transformed_sequence = transform_sequence(seq, groups, group_codes)
            #Only creates plots if the user chose to:
            if plots != "off":
                dipep_plot1 = dipep_plot(seq, all_amino_acids)
                pio.write_image(dipep_plot1, output + '_20_letters.png')
                dipep_plot2 = dipep_plot(transformed_sequence, amino_acids, groups)
                pio.write_image(dipep_plot2, output + "_" + str(len(amino_acids)) + '_letters.png')

        ## ======================================================================== TOP - LD  1/8/2021

            #prints protein/IDR name
            #print(name)


            #Creates pie chart of top freq amino acids
            #Only creates plots if the user chose to, but frequency calculations are still conducted inside the function:
            pie_plot1 = pie_chart(output, df, 'aa', 'freq', plots, name)
            #pie_plot1.savefig(str(OUTPUT)+'_pie_chart.png')
            #pio.write_image(pie_plot1, output + '_pie_chart.png')    

            #Creates amino acid frequency table and plots for modified sequence
            df2 = create_aa_df(transformed_sequence, amino_acids)
            #Only creates plots if the user chose to:
            if plots != "off":
                freq_plot3 = freq_plot(df2, 'aa', 'freq', name)
                pio.write_image(freq_plot3, output + '_freq_plot_codes.png')


            #Provides the dipeptide values for the text file Table
            dipep_values1 = dipep_plot_values(seq, all_amino_acids)
            dipep_values2 = dipep_plot_values(transformed_sequence, amino_acids, groups)
            #print(dipep_values1)
            #print(dipep_values2)


            #Provides the tripeptide values for the text file Table --- Modified on 1/28/2021
            tripep_values1 = tripep_plot_values(seq, all_amino_acids)
            tripep_values2 = tripep_plot_values(transformed_sequence, amino_acids, groups)
            #print(tripep_values1)
            #print(tripep_values2)


            #Converting data from graph to table
            #Creating handy dictionary
            df_both = {'df_1':[], 'df_2':[], 'df_1_name':[], 'df_2_name':[]}
            #Adding data from Csaba's df
            for element1 in df['freq']:  #For amino acids
                df_both['df_1'].append(element1)
            #print(str(df_both['df_1']))
            #Adding data from df2
            for element2 in df2['freq']:  #For coded/clustered version
                df_both['df_2'].append(element2)
            #print(str(df_both['df_2']))
            #Adding data from Csaba's df
            for element11 in df['aa']:  #For amino acids names
                df_both['df_1_name'].append(element11)
            #print(str(df_both['df_1']))
            #Adding data from df2
            for element12 in df2['aa']:  #For coded/clustered version names
                df_both['df_2_name'].append(element12)
            #print(str(df_both['df_2']))
            

            #Converting dipeptide data from graph to table
            #Creating handy dictionary
            dipep_both = {} #{'dipep_1':[], 'dipep_2':[]}
            dipep_both['dipep_1'] = {'dipeptide_name':[], 'freq':[]}  #For amino acids
            dipep_both['dipep_2'] = {'dipeptide_name':[], 'freq':[]}  #For coded/clustered version

            #Converting dipeptide data from graph to table --- Modified on 1/28/2021
            #Creating handy dictionary
            tripep_both = {} #{'tripep_1':[], 'tripep_2':[]}
            tripep_both['tripep_1'] = {'tripeptide_name':[], 'freq':[]}  #For amino acids
            tripep_both['tripep_2'] = {'tripeptide_name':[], 'freq':[]}  #For coded/clustered version


            #Adding dipeptide name data from dipep_values1
            for element3 in dipep_values1['dipeptide']:  #For amino acids
                dipep_both['dipep_1']['dipeptide_name'].append(element3)
            #Adding dipeptide name data from dipep_values2
            for element4 in dipep_values2['dipeptide']:  #For coded/clustered version
                dipep_both['dipep_2']['dipeptide_name'].append(element4)
            #Adding frequency data from dipep_values1
            for element5 in dipep_values1['freq']:  #For amino acids
                dipep_both['dipep_1']['freq'].append(element5)
            #Adding frequency data from dipep_values2
            for element6 in dipep_values2['freq']:  #For coded/clustered version
                dipep_both['dipep_2']['freq'].append(element6)

            #print(dipep_both['dipep_1'])
            #print(dipep_both['dipep_2'])


            #Adding tripeptide name data from tripep_values1 --- Modified on 1/28/2021
            for element7 in tripep_values1['tripeptide']:  #For amino acids
                tripep_both['tripep_1']['tripeptide_name'].append(element7)
            #Adding tripeptide name data from tripep_values2
            for element8 in tripep_values2['tripeptide']:  #For coded/clustered version
                tripep_both['tripep_2']['tripeptide_name'].append(element8)
            #Adding frequency data from tripep_values1
            for element9 in tripep_values1['freq']:  #For amino acids
                tripep_both['tripep_1']['freq'].append(element9)
            #Adding frequency data from tripep_values2
            for element10 in tripep_values2['freq']:  #For coded/clustered version
                tripep_both['tripep_2']['freq'].append(element10)

            #print(tripep_both['tripep_1'])
            #print(tripep_both['tripep_2'])

            #Creates the header of the text file Table
            if table_counter == 0:        
                f2 = open(str(output_dir)+"Table.txt", "a")
                f2.write("Name"+' - '+"Length"+' - ' + "Top1,Freq" +' '+ "Top2,Freq" +' '+ "Top3,Freq" +' '+ "Others" + ' - ' +str(all_amino_acids).strip("[]")+' - '+str(amino_acids).strip("[]")+' - '+str(dipep_both['dipep_2']['dipeptide_name']).strip("[]")+' - '+"Name"+"\n")
                f2.close()
                table_counter = table_counter + 1

            #Stores information in the text file Table
            f2 = open(str(output_dir)+"Table.txt", "a")
            f2.write(str(name)+' - '+str(len(seq))+' - '+ str(pie_plot1['labels'][0]).strip(">'[]'")+','+str(pie_plot1['freq'][0]).strip(">'[]'")+' '+ str(pie_plot1['labels'][1]).strip(">'[]'")+','+str(pie_plot1['freq'][1]).strip(">'[]'")+' '+ str(pie_plot1['labels'][2]).strip(">'[]'")+','+str(pie_plot1['freq'][2]).strip(">'[]'")+' '+str(pie_plot1['freq'][3]).strip(">'[]'") + ' - ' + str(df_both['df_1']).strip(">'[]'")+' - '+str(df_both['df_2']).strip(">'[]'")+' - '+str(dipep_both['dipep_2']['freq']).strip(">'[]'")+' - '+str(name)+"\n")
            #f2.write("Name"+' '+str(len(seq))+' '+str(all_amino_acids).strip("[]")+' '+str(amino_acids).strip("[]"))
            f2.close()

            #Creates the header of the extended text file Table  --- Modified on 1/28/2021
            if table_counter_extended == 0:        
                f3 = open(str(output_dir)+"Extended_Table.txt", "a")
                f3.write("Name"+' - '+"Length"+' - ' + "Top1,Freq" +' '+ "Top2,Freq" +' '+ "Top3,Freq" +' '+ "Others" + ' - ' +str(all_amino_acids).strip("[]")+' - '+str(amino_acids).strip("[]")+' - '+"Name"+' - '+str(dipep_both['dipep_2']['dipeptide_name']).strip("[]")+' - '+"Name"+' - '+str(tripep_both['tripep_2']['tripeptide_name']).strip("[]")+' - '+"Name"+' - '+str(dipep_both['dipep_1']['dipeptide_name']).strip("[]")+' - '+"Name"+' - '+str(tripep_both['tripep_1']['tripeptide_name']).strip("[]")+' - '+"Name"+"\n")
                f3.close()
                
            #Stores information in the extended text file Table  --- Modified on 1/28/2021
            f3 = open(str(output_dir)+"Extended_Table.txt", "a")
            f3.write(str(name)+' - '+str(len(seq))+' - '+ str(pie_plot1['labels'][0]).strip(">'[]'")+','+str(pie_plot1['freq'][0]).strip(">'[]'")+' '+ str(pie_plot1['labels'][1]).strip(">'[]'")+','+str(pie_plot1['freq'][1]).strip(">'[]'")+' '+ str(pie_plot1['labels'][2]).strip(">'[]'")+','+str(pie_plot1['freq'][2]).strip(">'[]'")+' '+str(pie_plot1['freq'][3]).strip(">'[]'") + ' - ' + str(df_both['df_1']).strip(">'[]'")+' - '+str(df_both['df_2']).strip(">'[]'")+' - '+str(name)+' - '+str(dipep_both['dipep_2']['freq']).strip(">'[]'")+' - '+str(name)+' - '+str(tripep_both['tripep_2']['freq']).strip(">'[]'")+' - '+str(name)+' - '+str(dipep_both['dipep_1']['freq']).strip(">'[]'")+' - '+str(name)+' - '+str(tripep_both['tripep_1']['freq']).strip(">'[]'")+' - '+str(name)+"\n")
            #f3.write("Name"+' '+str(len(seq))+' '+str(all_amino_acids).strip("[]")+' '+str(amino_acids).strip("[]"))
            f3.close()
            
        ## ======================================================================== Bottom - LD  1/8/2021    
            
            #Prepares to export the data in a database format in a dett file
            #Creates database dett file for current group and the header
            if table_counter_extended == 0:
                f11 = open(str(output_database)+ "Database_" + run_name +".dett", "a")
                #Writes all amino acids
                f11.write("pepname<" + str(all_amino_acids).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes all amino acid codes
                f11.write("codepepname<" + str(amino_acids).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes all dipeptides
                f11.write("dipepname<" + str(dipep_both['dipep_1']['dipeptide_name']).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes all dipeptide codes
                f11.write("codedipepname<" + str(dipep_both['dipep_2']['dipeptide_name']).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes all tripeptides
                f11.write("tripepname<" + str(tripep_both['tripep_1']['tripeptide_name']).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes all tripeptide codes
                f11.write("codetripepname<" + str(tripep_both['tripep_2']['tripeptide_name']).strip("[]").replace("'","").replace(" ","") + "\n")
                #Writes a blank line
                f11.write("\n")
                f11.close()
                table_counter_extended = table_counter_extended + 1
                
                
            #Once database dett file was created, starts to write the information of the current protein into the dett file
            f11 = open(str(output_database)+ "Database_" + run_name +".dett", "a")
            #Writes protein name
            f11.write("name<" + str(name).replace("'","").replace(" ","") + "\n")
            #Writes protein/truncation sequence
            f11.write("seq<" + str(seq).replace("'","").replace(" ","") + "\n")
            #Writes all amino acids
            f11.write("pepfreq<" + str(df_both['df_1']).strip("[]").replace("'","").replace(" ","") + "\n")
            #Writes all amino acid codes
            f11.write("codepepfreq<" + str(df_both['df_2']).strip("[]").replace("'","").replace(" ","") + "\n")
            #Writes all dipeptides
            f11.write("dipepfreq<" + str(dipep_both['dipep_1']['freq']).strip("[]").replace("'","").replace(" ","") + "\n")
            #Writes all dipeptide codes
            f11.write("codedipepfreq<" + str(dipep_both['dipep_2']['freq']).strip("[]").replace("'","").replace(" ","") + "\n")
            #Writes all tripeptides
            f11.write("tripepfreq<" + str(tripep_both['tripep_1']['freq']).strip("[]").replace("'","").replace(" ","") + "\n")
            #Writes all tripeptide codes
            f11.write("codetripepfreq<" + str(tripep_both['tripep_2']['freq']).strip("[]").replace("'","").replace(" ","") + "\n")
            
            

            #Writes a blank line
            f11.write("\n")
            f11.close()
                

        
        ## ======================================================================== Bottom - LD  7/30/2021
        
        ## ======================================================================== TOP - LD  2/25/2021

            
            #Starts to organize data using pandas to calculate top frequencies
            if counter == 0:  #First protein sequence/ first time
                #Updates dipeptide amino acid data
                s_avg_dipep_aa = pd.Series(dipep_both['dipep_1']['freq'], index=dipep_both['dipep_1']['dipeptide_name'])
                #Updates dipeptide clustered data
                s_avg_dipep_cluster = pd.Series(dipep_both['dipep_2']['freq'], index=dipep_both['dipep_2']['dipeptide_name'])
                #Updates tripeptide amino acid data
                s_avg_tripep_aa = pd.Series(tripep_both['tripep_1']['freq'], index=tripep_both['tripep_1']['tripeptide_name'])
                #Updates tripeptide clustered data
                s_avg_tripep_cluster = pd.Series(tripep_both['tripep_2']['freq'], index=tripep_both['tripep_2']['tripeptide_name'])
                #Updates single amino acid data
                s_avg_pep_aa = pd.Series(df_both['df_1'], index=df_both['df_1_name'])
                #Updates single peptide clustered data
                s_avg_pep_cluster = pd.Series(df_both['df_2'], index=df_both['df_2_name'])
                
            else:
                #Updates dipeptide amino acid data
                s_current_dipep_aa = pd.Series(dipep_both['dipep_1']['freq'], index=dipep_both['dipep_1']['dipeptide_name'])
                s_avg_dipep_aa = s_avg_dipep_aa + s_current_dipep_aa
                #Updates dipeptide clustered data
                s_current_dipep_cluster = pd.Series(dipep_both['dipep_2']['freq'], index=dipep_both['dipep_2']['dipeptide_name'])
                s_avg_dipep_cluster = s_avg_dipep_cluster + s_current_dipep_cluster
                #Updates tripeptide amino acid data
                s_current_tripep_aa = pd.Series(tripep_both['tripep_1']['freq'], index=tripep_both['tripep_1']['tripeptide_name'])
                s_avg_tripep_aa = s_avg_tripep_aa + s_current_tripep_aa
                #Updates tripeptide clustered data
                s_current_tripep_cluster = pd.Series(tripep_both['tripep_2']['freq'], index=tripep_both['tripep_2']['tripeptide_name'])
                s_avg_tripep_cluster = s_avg_tripep_cluster + s_current_tripep_cluster
                #Updates single amino acid data
                s_current_pep_aa = pd.Series(df_both['df_1'], index=df_both['df_1_name'])
                s_avg_pep_aa = s_avg_pep_aa + s_current_pep_aa
                #Updates single peptide clustered data
                s_current_pep_cluster = pd.Series(df_both['df_2'], index=df_both['df_2_name'])
                s_avg_pep_cluster = s_avg_pep_cluster + s_current_pep_cluster
                            
                
            #Updates neat counter
            counter = counter + 1
            

        ## ======================================================================== Bottom - LD  2/25/2021

            #Color-code and save the transformed sequence into a Word document
            #Only does so if the user chose to:
            if plots != "off":
                df = pd.DataFrame(list(transformed_sequence), columns=['letters'])
                p = document.add_heading('{}'.format(str(count) + '_' + name))
                p = document.add_paragraph('')
                for k in df.index:
                    letter = df['letters'][k]
                    run = p.add_run(letter)
                    #how it works, for every letter in the sequence, if the letter is one of our "coded groups", color it
                    for index, aa in enumerate(group_codes):
                        if letter == aa:
                            color = color_dict[index]
                            run.font.highlight_color = eval('WD_COLOR_INDEX.{}'.format(color))
                #p = document.add_page_break()

                document.save(output_dir + 'Color_coded_sequences.docx')
        
        ## ======================================================================== TOP - LD  2/25/2021
                  
        #Calculates the average of the top frequencies with the aid of 'counter' and sorts in descending order
        #For dipeptide amino acid data
        s_avg_dipep_aa = (s_avg_dipep_aa/counter).sort_values(0, ascending = False)
        #For dipeptide clustered data
        s_avg_dipep_cluster = (s_avg_dipep_cluster/counter).sort_values(0, ascending = False)
        #For tripeptide amino acid data
        s_avg_tripep_aa = (s_avg_tripep_aa/counter).sort_values(0, ascending = False)
        #For tripeptide clustered data
        s_avg_tripep_cluster = (s_avg_tripep_cluster/counter).sort_values(0, ascending = False)
        #For single amino acid     
        s_avg_pep_aa = (s_avg_pep_aa/counter).sort_values(0, ascending = False)
        #For single peptide clustered          
        s_avg_pep_cluster = (s_avg_pep_cluster/counter).sort_values(0, ascending = False)
        
        #Adds the top frequency results to their respective .txt files
        #For Dipeptide amino acid
        f4 = open(str(output_dir)+"Top_dipeptide_aa_average.txt", "a")
        f4.write(str(s_avg_dipep_aa.to_string()))
        f4.close()
        #For Dipeptide clustered
        f5 = open(str(output_dir)+"Top_dipeptide_clustered_average.txt", "a")
        f5.write(str(s_avg_dipep_cluster.to_string()))
        f5.close()
        #For Tripeptide amino acid
        f7 = open(str(output_dir)+"Top_tripeptide_aa_average.txt", "a")
        f7.write(str(s_avg_tripep_aa.to_string()))
        f7.close()
        #For Tripeptide clustered
        f8 = open(str(output_dir)+"Top_tripeptide_clustered_average.txt", "a")
        f8.write(str(s_avg_tripep_cluster.to_string()))
        f8.close()        
        #For single amino acid
        f9 = open(str(output_dir)+"Top_peptide_aa_average.txt", "a")
        f9.write(str(s_avg_pep_aa.to_string()))
        f9.close()
        #For single peptide clustered
        f10 = open(str(output_dir)+"Top_peptide_clustered_average.txt", "a")
        f10.write(str(s_avg_pep_cluster.to_string()))
        f10.close()        
        

        #Final message of this run
        print("\n" + 'Top Average Frequency txt files were successfully created!' + "\n")
        print(str(counter) + ' sequences were successfully processed!' + "\n")
        print("# --------------------------------------------- #")

    #Ending message
    print("\nAmino_Acid_Pattern_Finder_Extended is done!\n")
    print("# ------------------------------------------------------------------------------------------------------------------------ #")
    print("\n")    
    
    #Returns location of table_extended.txt ???  - to be continued
    #return df_both
